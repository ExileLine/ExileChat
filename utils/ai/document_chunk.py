# -*- coding: utf-8 -*-
# @Time    : 2024/7/27 16:50
# @Author  : yangyuexiong
# @Email   : yang6333yyx@126.com
# @File    : document_chunk.py
# @Software: PyCharm


import os
import base64
from io import BytesIO
from collections import OrderedDict

from docx import Document
from docx.oxml.ns import qn
from tabulate import tabulate


# import pandas as pd
# import pdfplumber


class DocumentChunk:
    """文档切片分块"""

    def __init__(self, image_base_path: str = None, image_base_url: str = None, is_debug: bool = False):
        self.image_base_path = image_base_path  # 磁盘存储目录路径
        self.image_base_url = image_base_url  # 图片服务器域名路径或静态资源路径
        self.is_debug = is_debug

        if not self.image_base_path and self.image_base_url:
            raise AttributeError("属性 image_base_path 与 image_base_url 必须要其中一个")

    def read_pdf(self, file_content):
        content = ""
        with pdfplumber.open(BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                content += page.extract_text() + "\n"
        return content

    def read_excel(self, file_content):
        df = pd.read_excel(BytesIO(file_content))
        return df.to_string()

    def read_docx(self, file_path):
        doc = Document(file_path)

        paragraph_dict = {}
        image_dict = {}
        table_dict = {}

        paragraph_index = []
        table_index = []
        image_counter = 1

        for index, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 处理段落
                if index in paragraph_index:
                    set_paragraph_index = index + 1
                else:
                    set_paragraph_index = index
                paragraph_dict[set_paragraph_index] = {"type": "paragraph", "content": paragraph.text}
                paragraph_index.append(set_paragraph_index)
                # print(set_paragraph_index, paragraph.text, paragraph_index)

            # 检查图片
            for run in paragraph.runs:
                for drawing in run.element.findall('.//w:drawing', namespaces=run.element.nsmap):
                    if paragraph_index:
                        set_image_index = max(paragraph_index) + 1
                    else:
                        set_image_index = index + 1

                    blip = drawing.find(
                        './/a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
                    )
                    if blip is not None:
                        rId = blip.get(qn('r:embed'))
                        image_part = doc.part.related_parts[rId]
                        image_bytes = image_part.blob
                        image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                        if self.image_base_path:
                            image_filename = self.image_base_path + f"/docx_image_{image_counter}.png"
                            with open(image_filename, 'wb') as image_file:
                                image_file.write(image_bytes)
                        elif self.image_base_url:
                            # TODO 补充上传图片服务逻辑
                            image_filename = self.image_base_url + f"/docx_image_{image_counter}.png"
                        else:
                            raise AttributeError("属性 image_base_path 与 image_base_url 必须要其中一个")

                        print(image_filename)
                        image_dict[set_image_index] = {"type": "image", "content": image_filename}
                        paragraph_index.append(set_image_index)
                        image_counter += 1

                    # print(set_image_index, paragraph_index)
                    # image_dict[set_image_index] = {"type": "image", "content": "Image found"}
                    # paragraph_index.append(set_image_index)

            if index not in paragraph_index:
                table_index.append(index)

        # print(table_index)

        # 处理表格
        for index, table in enumerate(doc.tables):
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text.strip())
                table_content.append(row_content)
            set_table_index = table_index[index]
            content = tabulate(table_content, headers="firstrow", tablefmt="grid")
            table_dict[set_table_index] = {"type": "table", "content": content}

        data = {**paragraph_dict, **image_dict, **table_dict}
        sorted_data = OrderedDict(sorted(data.items(), key=lambda item: int(item[0])))

        if self.is_debug:
            # 打印排序后的字典
            for key, value in sorted_data.items():
                print(key, value)

        result = self.generate_content(sorted_data)
        print(result)

    @staticmethod
    def generate_content(sorted_data: dict) -> str:
        """生成全文字符串"""

        result = ""
        for key, value in sorted_data.items():
            data_type = value.get("type")
            data_content = value.get("content")
            if data_type == "paragraph":
                result += f"{data_content}\n"
            elif data_type == "table":
                result += f"{data_content}\n"
            elif data_type == "image":
                result += f"{data_content}\n"
            else:
                result += f"{data_content}\n"
        return result

    def process_file(self, file_path):
        """主函数"""

        file_ext = os.path.splitext(file_path)[1].lower()
        print(file_ext)

        if file_ext == ".docx":
            self.read_docx(file_path)
        elif file_ext in [".xls", ".xlsx"]:
            self.read_excel(file_path)
        elif file_ext == ".pdf":
            self.read_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
